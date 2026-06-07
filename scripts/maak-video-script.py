# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GRAY = RGBColor(0x80, 0x80, 0x80)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)

def run(p, t, bold=False, italic=False, color=None, size=None):
    r = p.add_run(t); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    if size is not None: r.font.size = Pt(size)
    return r

p = doc.add_paragraph(); run(p, "Khalaf BI — Voice-over script schermvideo", bold=True, size=22, color=NAVY)
p = doc.add_paragraph(); run(p, "Demo Attiva Zorg · richttijd ± 5 minuten", bold=True, size=13, color=GOLD)
p = doc.add_paragraph(); run(p, "Gebruik khalaf-bi.vercel.app voor de opname", size=10, color=GRAY)

ph = doc.add_heading("Tips vooraf", level=1)
for r in ph.runs: r.font.color.rgb = NAVY
for b in [
    "Zet de tabbladen vooraf klaar: dashboard (ingelogd), de maandrapport-preview, en eventueel Supabase.",
    "Spreek rustig (±130 woorden per minuut) en klik kalm — geef de kijker tijd om mee te kijken.",
    "Opnemen kan met Windows Game Bar (Win+G), of met OBS / Loom.",
    "Log in met de docent-login of je eigen account; begin op het Attiva-dashboard.",
    "De cursief gedrukte regels ('Op het scherm') zijn aanwijzingen voor jezelf — die spreek je niet uit.",
]:
    doc.add_paragraph(b, style="List Bullet")

scenes = [
    ("Scène 1 — Intro & jouw 'why'", "0:00 – 0:35", "Homepage / login van khalaf-bi.vercel.app.",
     "Hoi, ik ben Adan, en dit is Khalaf BI — een dashboardportaal dat ik heb gebouwd voor het MKB. "
     "Mijn 'why' is simpel: ik zie dat ondernemers verdrinken in losse Excel-bestanden en pas veel te laat "
     "weten hoe hun bedrijf ervoor staat. Dat wilde ik oplossen — niet met nóg een spreadsheet, maar met één "
     "helder portaal dat automatisch de cijfers ophaalt en je vertelt wáár je op moet sturen. Ik laat het zien "
     "aan de hand van mijn klant, Attiva Zorg."),

    ("Scène 2 — Hoe ik het heb gebouwd & inloggen", "0:35 – 1:10", "Inloggen; je landt op het Attiva-dashboard.",
     "Technisch heb ik dit gebouwd met Next.js voor de website, Supabase als database, en een koppeling met "
     "Exact Online voor de financiële data — alles staat live online. Ik log nu in… en je komt meteen op het "
     "dashboard van Attiva Zorg. Bovenaan zie je direct de kern: de klantnaam, de koppelstatus met Exact, en een "
     "gezondheidsscore — een stoplicht dat in één oogopslag laat zien hoe het bedrijf ervoor staat."),

    ("Scène 3 — Wat vraagt aandacht?", "1:10 – 1:35", "Het blok 'Wat vraagt aandacht?'.",
     "Dit onderdeel, 'Wat vraagt aandacht', is mijn favoriet. In plaats van de ondernemer te laten zoeken, zet "
     "het systeem de belangrijkste signalen bovenaan — telkens mét een concrete actie erbij. Denk aan een negatief "
     "resultaat, of crediteuren die te lang openstaan. Dit is waar je je dag mee begint."),

    ("Scène 4 — Kerncijfers & 'Waarom?'", "1:35 – 2:10", "Management­samenvatting; klik op een 'Waarom?'-knop.",
     "Daaronder de kerncijfers: omzet, kosten, resultaat en marge, met de vergelijking ten opzichte van vorig jaar. "
     "En het mooie is: bij elk cijfer kun je doorklikken op 'Waarom?'. Dan laat het systeem zien hoe dat cijfer is "
     "opgebouwd — welke kostenposten stegen, welke omzet daalde. Zo blijft het niet bij een getal, maar snap je het "
     "verhaal erachter."),

    ("Scène 5 — Wat ik heb bereikt: tot 2026 & prognose", "2:10 – 2:50", "Wissel van jaar 2024 → 2025 → 2026; toon de Omzetvergelijking met prognose.",
     "Wat ik heb bereikt: het loopt niet alleen tot 2025, maar ook tot en met 2026. Ik wissel even van jaar… en hier, "
     "in de omzetvergelijking, zie je naast de gerealiseerde maanden ook een prognose voor de maanden die nog komen. "
     "En heel bewust heb ik de bankmutaties boekhoudkundig opgeschoond: privé-leningen en interne overboekingen tel ik "
     "níét mee. Daardoor klopt het resultaat — een realistisch beeld, geen mooiweer-cijfers."),

    ("Scène 6 — Grafieken & verdieping", "2:50 – 3:20", "Hero-grafiek Omzet vs Kosten, kosten-donut, leveranciers; tabbladen Cashflow/Crediteuren/Declaraties.",
     "De cijfers kun je ook visueel verkennen: omzet tegenover kosten per maand, de kostenverdeling, en de grootste "
     "leveranciers waaraan is betaald. En via de tabbladen kom je bij de cashflow, de crediteuren en het "
     "declaratieoverzicht — allemaal vanuit één bron."),

    ("Scène 7 — AI-assistent & presentatieweergave", "3:20 – 4:05",
     "Klik op 'Vraag je data' en stel een vraag (bijv. 'Wat waren mijn grootste kostenposten?'); toon het antwoord. Klik daarna rechtsboven op 'Acties' → Presentatieweergave (en wijs de PDF-rapport-optie aan).",
     "Een onderdeel waar ik echt trots op ben, is de ingebouwde AI-assistent. Via de knop 'Vraag je data' kan de "
     "ondernemer gewoon een vraag stellen in normale taal — bijvoorbeeld 'wat waren mijn grootste kostenposten?' — en "
     "krijgt meteen antwoord op basis van de eigen cijfers. Geen Excel-formules meer, gewoon vragen. En rechtsboven, "
     "onder het 'Acties'-menu, zit een presentatieweergave: met één klik krijg je een strakke, schermvullende weergave "
     "om de cijfers aan een klant te presenteren — plus de optie om er een PDF-rapport van te exporteren."),

    ("Scène 8 — Automatische maandrapportage & architectuur", "4:05 – 4:40", "Maandrapport-mail (preview); kort Supabase/Exact tonen.",
     "Het systeem werkt ook proactief: elke maand stuurt het automatisch een managementrapport per e-mail naar de klant, "
     "met de gezondheidsscore, de cijfers en de aandachtspunten. En dit alles draait op één geïntegreerde omgeving: de "
     "data komt beveiligd binnen uit Exact Online, wordt opgeslagen in Supabase, en is per gebruiker afgeschermd."),

    ("Scène 9 — Uitdagingen & afsluiter", "4:40 – 5:10", "Rustig terug naar het dashboard-overzicht / logo.",
     "Waar ik tegenaan liep? Vooral data-consistentie en privacy — als twee cijfers niet matchen, ben je het vertrouwen "
     "meteen kwijt. Daar heb ik veel tijd in gestoken: één bron, kloppende cijfers. En hoe ik hiernaar kijk: ik ben er "
     "trots op dat dit een écht werkend product is, live, op mijn eigen domein. Wat ik wil dat je onthoudt: Khalaf BI "
     "maakt cijfers begrijpelijk en stuurt op actie — niet op spreadsheets. Dank je wel."),
]

VRAGEN = [
    "Wat waren mijn grootste kostenposten dit jaar?",
    "Hoe ontwikkelt mijn omzet zich ten opzichte van vorig jaar?",
    "Waarom staat mijn resultaat onder druk?",
    "Welke kostenpost is het hardst gestegen?",
    "Hoeveel hebben we uitgegeven aan personeel en lonen?",
    "Welke leveranciers kregen het meeste betaald?",
    "Wat is mijn brutomarge en is die verbeterd ten opzichte van vorig jaar?",
    "Wat is de verwachte jaaromzet en het verwachte resultaat?",
    "Welke maand had het beste en welke het slechtste resultaat?",
]

for titel, tijd, beeld, vo in scenes:
    h = doc.add_heading(titel + "   (" + tijd + ")", level=2)
    for r in h.runs: r.font.color.rgb = NAVY
    p = doc.add_paragraph(); run(p, "Op het scherm: ", bold=True, italic=True, color=GRAY, size=10); run(p, beeld, italic=True, color=GRAY, size=10)
    p = doc.add_paragraph(); run(p, "Voice-over: ", bold=True, color=GOLD, size=11); run(p, vo, size=12)
    if "AI-assistent" in titel:
        pq = doc.add_paragraph(); run(pq, "Voorbeeldvragen die je aan de AI-assistent kunt stellen:", bold=True, color=NAVY, size=11)
        for q in VRAGEN:
            doc.add_paragraph(q, style="List Bullet")
    doc.add_paragraph()

woorden = sum(len(vo.split()) for *_ , vo in scenes)
p = doc.add_paragraph(); run(p, f"Totaal voice-over: ± {woorden} woorden (≈ {round(woorden/130)} min bij rustig spreektempo).", italic=True, color=GRAY, size=10)

out = r"C:\Users\adama\OneDrive - Windesheim Office365\Finance and Control jaar 4\Minor\Khalaf BI - Voice-over script schermvideo.docx"
doc.save(out)
print("OK ->", out, "| scenes:", len(scenes), "| woorden:", woorden)
