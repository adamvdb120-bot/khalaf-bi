# -*- coding: utf-8 -*-
import os, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MINOR = r"C:\Users\adama\OneDrive - Windesheim Office365\Finance and Control jaar 4\Minor"
SHOTS = {
    1: None,  # tabellenlijst — geen aparte afbeelding; lijst is links zichtbaar in de andere screenshots
    2: MINOR + r"\Exact data cache.png",
    3: MINOR + r"\Exact tokens.png",
    4: MINOR + r"\Profiles.png",
    5: MINOR + r"\Authentication.png",
    6: MINOR + r"\Client maandafsluiting.png",
    7: r"C:\Users\adama\Downloads\khalaf-screenshots\7_managementrapport.png",
}

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GRAY = RGBColor(0x80, 0x80, 0x80)

doc = Document()

# standaard lettertype
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def para(text, italic=False, color=None, size=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    if color is not None: r.font.color.rgb = color
    if size is not None: r.font.size = Pt(size)
    return p

def screenshot(label, num):
    path = SHOTS.get(num)
    if path and os.path.exists(path):
        doc.add_picture(path, width=Inches(6.3))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = cap.add_run("Figuur: " + label)
        rr.italic = True; rr.font.color.rgb = GRAY; rr.font.size = Pt(9)
    elif path is None:
        pass  # bewust geen afbeelding (bv. tabellenlijst — zie tekst-tabel)
    else:
        p = doc.add_paragraph()
        r = p.add_run("[ Screenshot invoegen: " + label + " ]")
        r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
    return

def kv_table(rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(htxt)
        rr.bold = True; rr.font.color.rgb = NAVY; rr.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(val)
            rr.font.size = Pt(10)
    return t

# ---- Titel ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Khalaf BI — Technische onderbouwing")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
sub = doc.add_paragraph()
r = sub.add_run("Database, koppelingen en geautomatiseerde managementrapportage")
r.font.size = Pt(13); r.font.color.rgb = GOLD; r.bold = True
meta = doc.add_paragraph()
r = meta.add_run("Bijlage bij de rapportage · Attiva Zorg · juni 2026")
r.font.size = Pt(10); r.font.color.rgb = GRAY

doc.add_paragraph()
para("Deze bijlage licht de technische werking van het Khalaf BI-portaal toe aan de hand van "
     "schermafbeeldingen uit Supabase (de database en het toegangsbeheer), de Exact Online-koppeling "
     "en de geautomatiseerde managementrapportage per e-mail. Elk onderdeel wordt één voor één uitgelegd, "
     "zodat duidelijk is hoe de cijfers binnenkomen, worden opgeslagen en worden gerapporteerd.")

# ---- 1. Architectuur ----
h1("1. Architectuur in het kort")
para("Het portaal bestaat uit vier samenwerkende onderdelen:")
for b in [
    "Frontend (de website): gebouwd in Next.js en gehost op Vercel (khalaf-bi.vercel.app).",
    "Database & authenticatie: Supabase (een PostgreSQL-database) — hier staan alle gegevens en gebruikersaccounts.",
    "Externe databron: Exact Online (de financiële administratie), gekoppeld via een beveiligde API met OAuth-tokens.",
    "Geautomatiseerde maandrapportage: een e-mail die maandelijks automatisch wordt verstuurd (via Resend), aangestuurd door een geplande taak (Vercel cron).",
]:
    doc.add_paragraph(b, style="List Bullet")
para("Daarnaast is er een tijdelijke bankimport: een aparte databron voor perioden die nog niet volledig in "
     "Exact zijn geboekt. Exact blijft de primaire bron en neemt automatisch weer over zodra de periode is bijgewerkt.")

# ---- 2. Tabellenlijst ----
h1("2. De databasestructuur (Supabase – Table Editor)")
para("De database bevat de volgende tabellen — samen vormen ze het 'geheugen' van het portaal: financiële data, "
     "klant- en gebruikersgegevens, taken en rapporten. De volledige tabellenlijst is ook links zichtbaar in de "
     "schermafbeeldingen hierna.")
screenshot("Supabase Table Editor – tabellenlijst", 1)
kv_table([
    ["attiva_declaraties", "Declaratiegegevens per cliënt/zorgverlener — basis voor de omzet per cliënt."],
    ["attiva_doelen", "Doelstellingen en targets (bijvoorbeeld een omzetdoel) per jaar."],
    ["client_maandafsluiting", "Maandafsluiting met checklist — de administratiestatus per maand."],
    ["client_taken", "Takenlijst en opvolging van aandachtspunten."],
    ["clients", "Stamgegevens van de klantorganisaties in het portaal."],
    ["exact_data_cache", "Gecachte financiële data uit Exact Online (en de tijdelijke bankimport)."],
    ["exact_tokens", "Beveiligingssleutels (tokens) voor de Exact Online-koppeling."],
    ["profiles", "Gebruikersprofielen met rol en klantkoppeling — het toegangsbeheer."],
    ["reports", "Opgeslagen en gegenereerde rapporten."],
    ["uploads", "Door gebruikers geüploade bestanden."],
], ["Tabel", "Functie"])

# ---- 3. exact_data_cache ----
h1("3. exact_data_cache — het financiële geheugen")
para("Deze tabel bevat per regel een combinatie van client_name (de klant) en cache_key (welk soort gegevens), "
     "met de bijbehorende gegevens opgeslagen als JSON in de kolom data, plus updated_at (laatst bijgewerkt). "
     "In dit voorbeeld staan 28 records voor Attiva.")
screenshot("Supabase – tabel exact_data_cache", 2)
para("Voorbeelden van cache_key en wat ze bevatten:")
kv_table([
    ["2025-2024 / 2024-2023", "Financiële data van een jaar plus het vergelijkingsjaar (voor de YoY-vergelijking)."],
    ["bankimport-2026", "De tijdelijke bankimport voor 2026 — een aparte bron naast Exact."],
    ["narratief-2025 / insights-2025", "De automatisch gegenereerde AI-conclusies en signalen."],
    ["omzet-per-klant-2025-…", "Omzet uitgesplitst per categorie/financier."],
    ["transacties-per-tegenpartij-…", "Detailgegevens voor de drill-downs in het dashboard."],
], ["cache_key", "Inhoud"])
para("De kolom updated_at laat zien wanneer de data voor het laatst is ververst. Dit gebeurt automatisch via een "
     "nachtelijke geplande taak (cron), zodat het dashboard altijd actuele cijfers toont.")

# ---- 4. exact_tokens ----
h1("4. exact_tokens — de Exact Online-koppeling")
para("Deze tabel maakt de beveiligde koppeling met Exact Online mogelijk. Per klant worden de volgende "
     "gegevens bewaard:")
screenshot("Supabase – tabel exact_tokens", 3)
kv_table([
    ["access_token", "Tijdelijke sleutel waarmee het portaal data ophaalt uit Exact Online."],
    ["refresh_token", "Sleutel om automatisch een nieuwe access_token te verkrijgen als deze verloopt."],
    ["expires_at", "Tijdstip waarop de access_token verloopt."],
    ["division", "Het administratienummer van de klant in Exact Online (hier: 3901287)."],
], ["Kolom", "Betekenis"])
para("Doordat de access_token automatisch wordt ververst met de refresh_token, blijft de koppeling actief zonder "
     "dat er telkens handmatig opnieuw moet worden ingelogd. De tokens worden uitsluitend server-side gebruikt, "
     "dus nooit zichtbaar in de browser — dat is veilig.")

# ---- 5. Toegangsbeheer ----
h1("5. Toegangsbeheer: profiles en Authentication")
para("Het portaal werkt met persoonlijke accounts. Supabase Authentication beheert de inloggegevens; de tabel "
     "profiles bepaalt vervolgens wat iemand mag zien via een rol en een klantkoppeling (client_slug).")
screenshot("Supabase – tabel profiles", 4)
screenshot("Supabase – Authentication – Users", 5)
kv_table([
    ["adam.abdulahi2004@gmail.com", "admin", "Ziet alle klanten en dashboards (eigenaar/beheerder)."],
    ["finance@attivazorg.nl", "client (attiva)", "De klant Attiva Zorg — ziet alleen het eigen dashboard."],
    ["docent@khalaf-bi.nl", "client (attiva)", "Tijdelijke docent-login — komt direct op het Attiva-dashboard."],
], ["Gebruiker", "Rol", "Toegang"])
para("De toegangslogica is eenvoudig en veilig: een admin ziet alles, een client ziet uitsluitend het dashboard "
     "van de eigen organisatie. Zo kan een klant nooit bij de gegevens van een andere klant.")

# ---- 6. client_maandafsluiting ----
h1("6. client_maandafsluiting — administratiestatus")
para("Deze tabel houdt per maand bij hoe ver de administratie is. Elke regel heeft een status "
     "(open, in verwerking, gecontroleerd of afgesloten) en een reeks controlevelden (check_boekingen, "
     "check_omzet, check_kosten, enzovoort).")
screenshot("Supabase – tabel client_maandafsluiting", 6)
para("Deze gegevens voeden het blok 'Administratiestatus' op het dashboard, waar per maand zichtbaar is welke "
     "controles zijn afgevinkt en of de maand is afgesloten.")

# ---- 7. Managementrapport per mail ----
h1("7. Geautomatiseerde managementrapportage per e-mail")
para("Elke maand stuurt het systeem automatisch (via een Vercel cron op de 1e van de maand) een managementrapport "
     "per e-mail naar de klant. De afzender is Khalaf BI. Het rapport bevat een samenvatting van de belangrijkste "
     "cijfers en concrete aandachtspunten.")
screenshot("E-mail – managementrapport (Jaaroverzicht Attiva Zorg 2025)", 7)
para("Opbouw van de e-mail (zie het voorbeeld):")
for b in [
    "Klantgezondheid: een stoplicht-score (hier 'Actie nodig') met het aantal urgente meldingen en aandachtspunten.",
    "Kerncijfers: omzet, kosten en marge, met de verandering ten opzichte van vorig jaar.",
    "Totaaloverzicht: totale omzet, totale kosten en het resultaat over de periode.",
    "Aandachtspunten & acties: concrete signalen met een voorgestelde actie (bijv. negatief resultaat, openstaande crediteuren, cliënten zonder budget).",
    "Knop 'Bekijk volledig dashboard': directe link naar de live cijfers, grafieken en de AI-assistent.",
]:
    doc.add_paragraph(b, style="List Bullet")
para("Slimme terugval: als het lopende jaar nog onvoldoende data heeft, rapporteert de mail automatisch het "
     "laatste volledige jaar en benoemt dit expliciet (in het voorbeeld: 'Laatste beschikbare periode · 2025'). "
     "Zo bevat de mail nooit lege of misleidende cijfers.")
para("Veiligheid: er is een preview-modus die de mail toont zonder te versturen; het daadwerkelijk versturen is "
     "beveiligd met een geheime sleutel, zodat niemand ongevraagd een mail naar de klant kan laten sturen.")

# ---- 8. Conclusie ----
h1("8. Samengevat")
para("De schermafbeeldingen laten samen zien dat het portaal één geïntegreerde informatieomgeving is: de cijfers "
     "komen uit een centrale, beveiligde bron (Exact Online via tokens), worden opgeslagen en gecombineerd in "
     "Supabase, zijn afgeschermd met toegangsbeheer per gebruiker, en worden maandelijks automatisch gerapporteerd "
     "per e-mail. Niet langer losse Excel-bestanden, maar één betrouwbare bron die op actie stuurt.")

import os
out = r"C:\Users\adama\OneDrive - Windesheim Office365\Finance and Control jaar 4\Minor\Khalaf BI - Technische onderbouwing (database en rapportage).docx"
doc.save(out)
print("OK ->", out)
